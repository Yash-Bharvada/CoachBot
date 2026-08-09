"""Module 1a — Resume file parsing + structured Candidate Profile extraction.

Public surface:
  * :func:`extract_resume_text` — raw byte -> string, format-aware.
    Raises :class:`ResumeParsingError` (422) on unsupported extension,
    corrupted PDF/DOCX, or a file that exceeds the configured size cap.
  * :func:`parse_candidate_profile` — wraps :func:`extract_resume_text` and
    runs an LLM pass to return a structured :class:`CandidateProfile`.
  * :func:`build_candidate_highlights` — turns a CandidateProfile into the
    3-5 short strings surfaced on the :class:`OnboardResponse`.

Raw resume bytes are **never** persisted to Mongo.  The structured profile
lives inside ``role_context_matrices.candidate_profile``; the raw upload is
released from RAM as soon as parsing returns.
"""

from __future__ import annotations

import io
from typing import Any

from fastapi import UploadFile
from structlog import get_logger

from app.core.config import get_settings
from app.core.exceptions import ResumeParsingError
from app.models.schemas import CandidateProfile
from app.services.llm_client import get_groq_client

log = get_logger(__name__)

_RESUME_PARSE_SYSTEM = """You are a senior technical recruiter parsing a candidate resume.
Return STRICT JSON with these keys only (no extra prose):
{
  "skills": ["list of 8-16 specific, de-duplicated technical + soft skills mentioned or strongly implied"],
  "past_roles": ["Company, Job Title, Duration — one short line per significant role, most recent first, max 5"],
  "notable_projects": ["3-5 highest-signal projects; include scope, tech, outcome when stated; under 40 words each"],
  "education": ["Degrees, certifications, institutions with year when available; skip trivial entries"]
}
If a section has no signal return an empty [] for that key.  Do not invent content."""


# ---------------------------------------------------------------------------
# Raw text extraction (format-specific, no LLM)
# ---------------------------------------------------------------------------


def _extract_extension(filename: str) -> str:
    """Return a normalised lowercase extension (no leading dot)."""
    if "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].strip().lower()


async def _enforce_size_and_format(upload: UploadFile) -> bytes:
    """Read the upload body, enforcing extension + size caps.  Raises 422."""
    settings = get_settings()

    # --- 1. Format check (extension whitelist) -----------------------------
    ext = _extract_extension(upload.filename or "")
    if ext not in {e.lower() for e in settings.allowed_resume_extensions}:
        raise ResumeParsingError(
            "Unsupported resume file format.",
            details={
                "filename": upload.filename,
                "extension": ext or "(none)",
                "allowed": settings.allowed_resume_extensions,
            },
        )

    # --- 2. Read + size cap (stream-guard the in-memory buffer) ------------
    max_bytes = settings.max_resume_size_mb * 1024 * 1024
    chunks: list[bytes] = []
    total = 0
    while True:
        chunk = await upload.read(64 * 1024)
        if not chunk:
            break
        total += len(chunk)
        if total > max_bytes:
            raise ResumeParsingError(
                "Resume file exceeds the allowed size limit.",
                details={
                    "max_mb": settings.max_resume_size_mb,
                    "filename": upload.filename,
                },
            )
        chunks.append(chunk)
    raw = b"".join(chunks)
    if not raw:
        raise ResumeParsingError(
            "Uploaded resume file is empty.",
            details={"filename": upload.filename},
        )
    return raw


def _extract_pdf_text(raw: bytes, filename: str) -> str:
    """Extract text from a PDF using :mod:`pypdf`.  Raises 422 on corrupt data."""
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:  # pragma: no cover — enforced by requirements.txt
        raise ResumeParsingError(
            "PDF parsing dependency is not installed.",
            details={"filename": filename},
        ) from exc
    try:
        reader = PdfReader(io.BytesIO(raw))
        pages: list[str] = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001 — page-level defensive
                continue
        text = "\n".join(pages).strip()
    except Exception as exc:  # noqa: BLE001 — any pypdf-level corruption
        raise ResumeParsingError(
            "PDF file is corrupt or cannot be parsed.",
            details={"filename": filename, "error": type(exc).__name__},
        ) from exc
    if len(text) < 50:
        raise ResumeParsingError(
            "Resume PDF contained too little extractable text to parse.",
            details={"filename": filename, "extracted_chars": len(text)},
        )
    return text


def _extract_docx_text(raw: bytes, filename: str) -> str:
    """Extract text from a DOCX using :mod:`python-docx`.  Raises 422 on corrupt data."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover — enforced by requirements.txt
        raise ResumeParsingError(
            "DOCX parsing dependency is not installed.",
            details={"filename": filename},
        ) from exc
    try:
        doc = Document(io.BytesIO(raw))
        paragraphs: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())
        # Also pull text out of tables (common in one-column resumes).
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs).strip()
    except Exception as exc:  # noqa: BLE001 — any python-docx-level corruption
        raise ResumeParsingError(
            "DOCX file is corrupt or cannot be parsed.",
            details={"filename": filename, "error": type(exc).__name__},
        ) from exc
    if len(text) < 50:
        raise ResumeParsingError(
            "Resume DOCX contained too little extractable text to parse.",
            details={"filename": filename, "extracted_chars": len(text)},
        )
    return text


async def extract_resume_text(upload: UploadFile) -> tuple[str, str]:
    """Return (raw_text, extension) for the given upload.

    The caller is responsible for closing ``upload``; this function only
    consumes the body via ``await upload.read()`` so FastAPI cleans up the
    temporary file on response end.
    """
    raw = await _enforce_size_and_format(upload)
    filename = upload.filename or "resume"
    ext = _extract_extension(filename)
    if ext == "pdf":
        return _extract_pdf_text(raw, filename), ext
    if ext == "docx":
        return _extract_docx_text(raw, filename), ext
    # The _enforce_size_and_format gate already raised 422 for unknown
    # extensions; this branch is a defensive no-op to appease type checkers.
    raise ResumeParsingError(
        "Unsupported resume file format.",
        details={"filename": filename, "extension": ext},
    )


# ---------------------------------------------------------------------------
# LLM structured extraction
# ---------------------------------------------------------------------------


async def parse_candidate_profile(upload: UploadFile) -> CandidateProfile:
    """Full pipeline: validate, extract text, LLM -> CandidateProfile."""
    resume_text, ext = await extract_resume_text(upload)
    client = await get_groq_client()
    # Clip very long resumes to a window the LLM can handle reliably.
    clip_len = 12_000
    clipped = resume_text if len(resume_text) <= clip_len else resume_text[:clip_len]
    data: dict[str, Any] = await client.chat_completion_json(
        messages=[
            {"role": "system", "content": _RESUME_PARSE_SYSTEM},
            {
                "role": "user",
                "content": (
                    f"Resume file: {upload.filename or 'resume'} (format: {ext})\n\n"
                    f"Extracted resume text:\n{clipped}\n\n"
                    "Return valid JSON as instructed."
                ),
            },
        ],
        temperature=0.0,
        max_tokens=1500,
    )
    skills = [str(s).strip() for s in (data.get("skills") or []) if str(s).strip()][:20]
    past_roles = [str(r).strip() for r in (data.get("past_roles") or []) if str(r).strip()][:8]
    projects = [str(p).strip() for p in (data.get("notable_projects") or []) if str(p).strip()][:8]
    education = [str(e).strip() for e in (data.get("education") or []) if str(e).strip()][:8]
    return CandidateProfile(
        skills=skills,
        past_roles=past_roles,
        notable_projects=projects,
        education=education,
    )


# ---------------------------------------------------------------------------
# Frontend-facing highlights builder
# ---------------------------------------------------------------------------


def build_candidate_highlights(
    profile: CandidateProfile, max_items: int = 5
) -> list[str]:
    """Return a short, high-signal list for the onboarding confirmation preview.

    Ordering priority: most recent role (1) → strongest project(s) (2-3) →
    top skills that overlap with the JD (filled in by onboarding_service; here
    we just rank raw frequency/position).  Returns 3-5 items; never more
    than *max_items*.
    """
    highlights: list[str] = []
    if profile.past_roles:
        highlights.append(f"Most recent role: {profile.past_roles[0]}")
    if profile.notable_projects:
        for proj in profile.notable_projects[:2]:
            highlights.append(f"Key project: {proj}")
    if profile.skills:
        top_skills = profile.skills[: max(0, max_items - len(highlights))]
        if top_skills:
            highlights.append("Core skills: " + ", ".join(top_skills))
    if profile.education and len(highlights) < max_items:
        highlights.append(f"Education: {profile.education[0]}")
    return highlights[:max_items]
